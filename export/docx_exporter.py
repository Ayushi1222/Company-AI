import logging
from io import BytesIO
from typing import Dict
from datetime import datetime

logger = logging.getLogger(__name__)


class DOCXExporter:
    def __init__(self):
        self.docx_available = False
        
        try:
            from docx import Document
            from docx.shared import Inches, Pt, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            self.docx_available = True
            logger.info("python-docx available for DOCX export")
        except ImportError:
            logger.warning("python-docx not available")
    
    def export(self, plan: Dict, filename: str = None) -> BytesIO:
        if not self.docx_available:
            logger.error("python-docx not installed")
            raise ImportError("python-docx required for DOCX export")
        
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        doc = Document()
        
        title = doc.add_heading('ACCOUNT PLAN', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        company_name = plan['metadata']['company_name']
        company_heading = doc.add_heading(company_name, level=1)
        company_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        date_para = doc.add_paragraph(
            f"Generated: {datetime.now().strftime('%B %d, %Y')}"
        )
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_page_break()
        
        for section_name, section in plan['sections'].items():
            doc.add_heading(section['title'], level=1)
            
            content = section['content']
            if isinstance(content, dict):
                for key, value in content.items():
                    key_formatted = key.replace('_', ' ').title()
                    
                    p = doc.add_paragraph()
                    p.add_run(f"{key_formatted}: ").bold = True
                    
                    if isinstance(value, list):
                        doc.add_paragraph()
                        for item in value:
                            doc.add_paragraph(str(item), style='List Bullet')
                    else:
                        p.add_run(str(value))
                    
                    doc.add_paragraph()
        
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        return buffer
